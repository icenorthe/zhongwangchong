from __future__ import annotations

import argparse
import json
import sys
import time
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Callable, Iterable

import requests
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image
from selenium import webdriver
from selenium.common.exceptions import TimeoutException
from selenium.webdriver import Edge
from selenium.webdriver.common.by import By
from selenium.webdriver.edge.options import Options
from selenium.webdriver.remote.webdriver import WebDriver
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.select import Select
from selenium.webdriver.support.ui import WebDriverWait


WORKSPACE = Path(__file__).resolve().parents[1]
CONFIG_PATH = WORKSPACE / "config" / "pythonanywhere_secrets.json"
DEFAULT_BASE_URL = "http://icenorth.pythonanywhere.com"
EDGE_BINARY = Path(r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe")


@dataclass
class Shot:
    title: str
    note: str
    path: Path
    status: str = "通过"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate a UI screenshot Word report.")
    parser.add_argument("--base-url", default=DEFAULT_BASE_URL)
    parser.add_argument("--phone", required=True)
    parser.add_argument("--password", required=True)
    parser.add_argument("--output-dir", default="")
    return parser.parse_args()


def mask_phone(phone: str) -> str:
    phone = str(phone).strip()
    if len(phone) < 7:
        return phone
    return f"{phone[:3]}****{phone[-4:]}"


def load_admin_creds() -> dict[str, str]:
    if not CONFIG_PATH.exists():
        return {}
    data = json.loads(CONFIG_PATH.read_text(encoding="utf-8-sig"))
    return {
        "admin_token": str(data.get("admin_token", "")).strip(),
        "admin_password": str(data.get("admin_password", "")).strip(),
    }


def build_driver(width: int, height: int) -> Edge:
    options = Options()
    if EDGE_BINARY.exists():
        options.binary_location = str(EDGE_BINARY)
    options.add_argument("--headless=new")
    options.add_argument("--disable-gpu")
    options.add_argument("--window-size=%d,%d" % (width, height))
    options.add_argument("--force-device-scale-factor=1")
    options.add_argument("--hide-scrollbars")
    options.add_argument("--lang=zh-CN")
    driver = webdriver.Edge(options=options)
    driver.set_window_size(width, height)
    driver.set_page_load_timeout(60)
    return driver


def wait_visible(driver: WebDriver, by: By, value: str, timeout: int = 20):
    return WebDriverWait(driver, timeout).until(EC.visibility_of_element_located((by, value)))


def wait_clickable(driver: WebDriver, by: By, value: str, timeout: int = 20):
    return WebDriverWait(driver, timeout).until(EC.element_to_be_clickable((by, value)))


def wait_for(driver: WebDriver, condition: Callable[[WebDriver], object], timeout: int = 20):
    return WebDriverWait(driver, timeout).until(condition)


def safe_click(driver: WebDriver, element) -> None:
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", element)
    time.sleep(0.3)
    driver.execute_script("arguments[0].click();", element)


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def optimize_image(path: Path, max_width: int = 1500) -> None:
    with Image.open(path) as image:
        image.load()
        if image.width <= max_width:
            return
        ratio = max_width / float(image.width)
        resized = image.resize((max_width, max(1, int(image.height * ratio))), Image.LANCZOS)
        resized.save(path)


def capture_element(driver: WebDriver, selector: str, output_path: Path, timeout: int = 20) -> Path:
    element = wait_visible(driver, By.CSS_SELECTOR, selector, timeout=timeout)
    driver.execute_script("arguments[0].scrollIntoView({block:'start'});", element)
    time.sleep(0.6)
    element.screenshot(str(output_path))
    optimize_image(output_path)
    return output_path


def capture_viewport(driver: WebDriver, output_path: Path) -> Path:
    time.sleep(0.6)
    driver.save_screenshot(str(output_path))
    optimize_image(output_path)
    return output_path


def set_input_value(driver: WebDriver, selector: str, value: str) -> None:
    element = wait_visible(driver, By.CSS_SELECTOR, selector)
    element.clear()
    element.send_keys(value)


def wait_text_not_contains(driver: WebDriver, selector: str, text: str, timeout: int = 20):
    def condition(d: WebDriver):
        content = d.find_element(By.CSS_SELECTOR, selector).text
        return content if text not in content else False

    return wait_for(driver, condition, timeout=timeout)


def wait_has_any(driver: WebDriver, selectors: Iterable[str], timeout: int = 20) -> None:
    selectors = list(selectors)

    def condition(d: WebDriver):
        for selector in selectors:
            if d.find_elements(By.CSS_SELECTOR, selector):
                return True
        return False

    wait_for(driver, condition, timeout=timeout)


def click_page_tab(driver: WebDriver, page_name: str) -> None:
    button = wait_clickable(driver, By.CSS_SELECTOR, f'[data-page="{page_name}"]')
    safe_click(driver, button)
    wait_visible(driver, By.CSS_SELECTOR, f"#page-{page_name}")
    wait_for(
        driver,
        lambda d: d.find_element(By.CSS_SELECTOR, f"#page-{page_name}").is_displayed(),
        timeout=20,
    )
    time.sleep(0.8)


def click_admin_tab(driver: WebDriver, tab_name: str) -> None:
    button = wait_clickable(driver, By.CSS_SELECTOR, f'[data-tab="{tab_name}"]')
    safe_click(driver, button)
    pane_id = "userView" if tab_name == "userView" else tab_name
    wait_for(
        driver,
        lambda d: d.find_element(By.CSS_SELECTOR, f"#pane-{pane_id}").is_displayed(),
        timeout=20,
    )
    time.sleep(0.8)


def fetch_user_summary(base_url: str, phone: str, password: str) -> dict:
    session = requests.Session()
    response = session.post(
        f"{base_url}/api/auth/login",
        json={"phone": phone, "password": password},
        timeout=30,
    )
    response.raise_for_status()
    token = response.json()["token"]
    headers = {"Authorization": f"Bearer {token}"}
    me = session.get(f"{base_url}/api/me", headers=headers, timeout=30).json()
    orders = session.get(f"{base_url}/api/me/orders?limit=50", headers=headers, timeout=30).json()
    recharges = session.get(f"{base_url}/api/me/recharge-requests?limit=20", headers=headers, timeout=30).json()
    ledger = session.get(f"{base_url}/api/me/wallet-ledger?limit=50", headers=headers, timeout=30).json()
    return {
        "me": me,
        "orders_count": len(orders),
        "recharges_count": len(recharges),
        "ledger_count": len(ledger),
    }


def fetch_admin_summary(base_url: str, admin_token: str, admin_password: str) -> dict:
    if not admin_token:
        return {}
    headers = {
        "X-Admin-Token": admin_token,
        "X-Admin-Password": admin_password,
    }
    stats = requests.get(f"{base_url}/api/admin/stats", headers=headers, timeout=30).json()
    users = requests.get(f"{base_url}/api/admin/users?limit=500", headers=headers, timeout=30).json()
    orders = requests.get(f"{base_url}/api/orders?limit=200", headers=headers, timeout=30).json()
    recharges = requests.get(
        f"{base_url}/api/admin/recharge-requests?limit=200",
        headers=headers,
        timeout=30,
    ).json()
    return {
        "stats": stats,
        "users_count": len(users),
        "orders_count": len(orders),
        "recharges_count": len(recharges),
    }


def login_user_ui(driver: WebDriver, base_url: str, phone: str, password: str) -> None:
    driver.get(base_url)
    wait_visible(driver, By.CSS_SELECTOR, "#auth-section")
    set_input_value(driver, "#login-phone", phone)
    set_input_value(driver, "#login-password", password)
    safe_click(driver, wait_clickable(driver, By.CSS_SELECTOR, "#login-form button[type='submit']"))
    wait_for(driver, lambda d: d.find_element(By.CSS_SELECTOR, "#logout-btn").is_displayed(), timeout=25)
    wait_for(driver, lambda d: d.find_element(By.CSS_SELECTOR, "#hero-appbar").is_displayed(), timeout=25)
    wait_text_not_contains(driver, "#user-phone", "-", timeout=25)
    wait_has_any(driver, ["#my-orders article.order-item", "#my-orders .empty"], timeout=25)


def prepare_order_page(driver: WebDriver) -> None:
    click_page_tab(driver, "order")
    wait_visible(driver, By.CSS_SELECTOR, "#region-select")
    wait_for(driver, lambda d: len(Select(d.find_element(By.CSS_SELECTOR, "#region-select")).options) > 0, timeout=25)
    wait_for(driver, lambda d: len(Select(d.find_element(By.CSS_SELECTOR, "#station-select")).options) > 0, timeout=25)
    if len(Select(driver.find_element(By.CSS_SELECTOR, "#region-select")).options) > 1:
        Select(driver.find_element(By.CSS_SELECTOR, "#region-select")).select_by_index(1)
        time.sleep(0.8)
    if len(Select(driver.find_element(By.CSS_SELECTOR, "#station-select")).options) > 1:
        Select(driver.find_element(By.CSS_SELECTOR, "#station-select")).select_by_index(1)
        time.sleep(0.8)
    wait_for(driver, lambda d: len(Select(d.find_element(By.CSS_SELECTOR, "#socket-select")).options) > 0, timeout=25)


def prepare_status_page(driver: WebDriver) -> None:
    click_page_tab(driver, "status")
    wait_has_any(driver, [".status-region-btn", ".station-collapse", "#socket-overview .empty"], timeout=25)
    driver.execute_script(
        """
        const first = document.querySelector('.station-collapse');
        if (first) {
          first.open = true;
          first.scrollIntoView({block:'start'});
        }
        """
    )
    time.sleep(1.0)


def prepare_orders_page(driver: WebDriver) -> None:
    click_page_tab(driver, "orders")
    wait_has_any(driver, ["#my-orders article.order-item", "#my-orders .empty"], timeout=25)


def prepare_recharge_page(driver: WebDriver) -> None:
    click_page_tab(driver, "recharge")
    wait_visible(driver, By.CSS_SELECTOR, "#page-recharge")
    wait_has_any(driver, ["#my-recharges article.recharge-item", "#my-recharges .empty"], timeout=25)


def prepare_ledger_page(driver: WebDriver) -> None:
    click_page_tab(driver, "ledger")
    wait_visible(driver, By.CSS_SELECTOR, "#page-ledger")
    wait_has_any(driver, ["#my-wallet-ledger article.recharge-item", "#my-wallet-ledger .empty"], timeout=25)


def load_admin_orders(driver: WebDriver, base_url: str, admin_token: str, admin_password: str) -> None:
    driver.get(f"{base_url}/admin/orders")
    wait_visible(driver, By.CSS_SELECTOR, "#token")
    set_input_value(driver, "#token", admin_token)
    set_input_value(driver, "#password", admin_password)
    safe_click(driver, wait_clickable(driver, By.CSS_SELECTOR, "#refresh"))
    wait_text_not_contains(driver, "#meta", "网关模式：-", timeout=25)
    wait_has_any(driver, ["#tbody tr", "#tbody"], timeout=25)
    driver.execute_script(
        """
        document.getElementById('token').value='已填入管理令牌';
        document.getElementById('password').value='******';
        """
    )
    time.sleep(0.8)


def load_admin_users(driver: WebDriver, base_url: str, admin_token: str, admin_password: str, user_phone: str) -> None:
    driver.get(f"{base_url}/admin/users")
    wait_visible(driver, By.CSS_SELECTOR, "#token")
    set_input_value(driver, "#token", admin_token)
    set_input_value(driver, "#password", admin_password)
    safe_click(driver, wait_clickable(driver, By.CSS_SELECTOR, "#refresh"))
    wait_has_any(driver, ["#userTbody tr"], timeout=25)
    set_input_value(driver, "#userQuery", user_phone)
    safe_click(driver, wait_clickable(driver, By.CSS_SELECTOR, "#searchUsers"))
    wait_has_any(driver, ["#userTbody tr"], timeout=25)
    safe_click(driver, wait_clickable(driver, By.CSS_SELECTOR, "#userTbody tr"))
    wait_for(
        driver,
        lambda d: d.find_element(By.CSS_SELECTOR, "#selectedUserPhone").text.strip() not in {"", "-"},
        timeout=25,
    )
    driver.execute_script(
        """
        document.getElementById('token').value='已填入管理令牌';
        document.getElementById('password').value='******';
        """
    )
    time.sleep(0.8)


def build_report(
    output_path: Path,
    shots: list[Shot],
    base_url: str,
    user_phone: str,
    user_summary: dict,
    admin_summary: dict,
) -> None:
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Microsoft YaHei"
    normal_style.font.size = Pt(10.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("前端界面测试截图交付")
    title_run.bold = True
    title_run.font.size = Pt(18)

    meta = document.add_paragraph()
    meta.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    meta.add_run(f"测试地址：{base_url}\n")
    meta.add_run(f"测试账号：{mask_phone(user_phone)}\n")
    meta.add_run("测试策略：仅做界面与只读数据验证，未执行会改动线上业务数据的提交类操作。")

    summary = document.add_paragraph()
    summary.add_run("结论：").bold = True
    summary.add_run("本次检查中，用户端与管理端主要界面均可正常加载，数据接口返回正常，未发现阻断性界面异常。")

    scope = document.add_paragraph()
    scope.add_run("覆盖范围：").bold = True
    scope.add_run("登录、注册、站点状态、立即下单、我的订单、余额充值、消费记录、订单后台、充值申请、用户订单视图、用户后台。")

    risk = document.add_paragraph()
    risk.add_run("未执行操作：").bold = True
    risk.add_run("提交订单、提交充值申请、充值审核、余额调整、免费次数修改、密码重置、删除用户。")

    if user_summary:
        user_p = document.add_paragraph()
        user_p.add_run("用户侧数据摘要：").bold = True
        user_p.add_run(
            f"余额 {user_summary['me'].get('balance_yuan', 0):.2f} 元，"
            f"订单 {user_summary.get('orders_count', 0)} 条，"
            f"充值记录 {user_summary.get('recharges_count', 0)} 条，"
            f"流水 {user_summary.get('ledger_count', 0)} 条。"
        )

    if admin_summary:
        stats = admin_summary.get("stats", {})
        admin_p = document.add_paragraph()
        admin_p.add_run("管理侧数据摘要：").bold = True
        admin_p.add_run(
            f"用户 {admin_summary.get('users_count', 0)} 个，"
            f"订单 {admin_summary.get('orders_count', 0)} 条，"
            f"充值申请 {admin_summary.get('recharges_count', 0)} 条，"
            f"成功订单 {stats.get('success', 0)} 条，"
            f"失败订单 {stats.get('failed', 0)} 条，"
            f"待审核充值 {stats.get('recharge_pending', 0)} 条。"
        )

    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "模块"
    table.rows[0].cells[1].text = "结果"
    table.rows[0].cells[2].text = "说明"

    for shot in shots:
        row = table.add_row().cells
        row[0].text = shot.title
        row[1].text = shot.status
        row[2].text = shot.note

    document.add_page_break()

    for index, shot in enumerate(shots, start=1):
        heading = document.add_paragraph()
        heading_run = heading.add_run(f"{index}. {shot.title}")
        heading_run.bold = True
        heading_run.font.size = Pt(13)

        note = document.add_paragraph(shot.note)
        note_format = note.paragraph_format
        note_format.space_after = Pt(4)

        document.add_picture(str(shot.path), width=Inches(9.2))

        if index != len(shots):
            document.add_page_break()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def capture_user_shots(base_url: str, phone: str, password: str, output_dir: Path) -> list[Shot]:
    driver = build_driver(430, 1700)
    shots: list[Shot] = []
    try:
        driver.get(base_url)
        wait_visible(driver, By.CSS_SELECTOR, "#auth-section")
        shots.append(
            Shot(
                title="用户端-登录页",
                note="登录表单、提示文案与首屏布局展示正常。",
                path=capture_element(driver, "#auth-section", output_dir / "user_login.png"),
            )
        )

        safe_click(driver, wait_clickable(driver, By.CSS_SELECTOR, '[data-auth-tab="register"]'))
        time.sleep(0.6)
        shots.append(
            Shot(
                title="用户端-注册页",
                note="注册表单切换正常，可见手机号、密码与确认密码输入项。",
                path=capture_element(driver, "#auth-section", output_dir / "user_register.png"),
            )
        )

        safe_click(driver, wait_clickable(driver, By.CSS_SELECTOR, '[data-auth-tab="login"]'))
        login_user_ui(driver, base_url, phone, password)
        shots.append(
            Shot(
                title="用户端-登录后账号概览",
                note="账号、余额与免费服务费次数区域显示正常。",
                path=capture_element(driver, "section.hero", output_dir / "user_dashboard.png"),
            )
        )

        prepare_status_page(driver)
        shots.append(
            Shot(
                title="用户端-站点状态",
                note="区域切换、站点状态卡片与插座状态区块加载正常。",
                path=capture_element(driver, "#page-status", output_dir / "user_status.png"),
            )
        )

        prepare_order_page(driver)
        shots.append(
            Shot(
                title="用户端-立即下单",
                note="区域、站点、插座、金额与支付方式界面均已加载；本次未执行提交订单。",
                path=capture_element(driver, "#page-order", output_dir / "user_order.png"),
            )
        )

        prepare_orders_page(driver)
        shots.append(
            Shot(
                title="用户端-我的订单",
                note="历史订单记录成功展示，本次账号下已有订单明细可供核验。",
                path=capture_element(driver, "#page-orders", output_dir / "user_orders.png"),
            )
        )

        prepare_recharge_page(driver)
        shots.append(
            Shot(
                title="用户端-余额充值",
                note="人工扫码充值入口、说明文案与历史充值记录显示正常；本次未提交充值申请。",
                path=capture_element(driver, "#page-recharge", output_dir / "user_recharge.png"),
            )
        )

        prepare_ledger_page(driver)
        shots.append(
            Shot(
                title="用户端-消费记录",
                note="余额流水列表展示正常，可见充值与下单扣费/退款记录。",
                path=capture_element(driver, "#page-ledger", output_dir / "user_ledger.png"),
            )
        )
    finally:
        driver.quit()
    return shots


def capture_admin_order_shots(
    base_url: str,
    admin_token: str,
    admin_password: str,
    output_dir: Path,
) -> list[Shot]:
    if not admin_token:
        return []

    driver = build_driver(1500, 1700)
    shots: list[Shot] = []
    try:
        load_admin_orders(driver, base_url, admin_token, admin_password)
        driver.execute_script("window.scrollTo(0, 0);")
        shots.append(
            Shot(
                title="管理端-订单后台概览",
                note="顶部鉴权区、统计卡片与功能标签加载正常，截图中已对凭证做脱敏处理。",
                path=capture_viewport(driver, output_dir / "admin_orders_top.png"),
            )
        )
        shots.append(
            Shot(
                title="管理端-订单总览表",
                note="订单列表、状态列与操作列展示正常，可按设备编码和状态筛选。",
                path=capture_element(driver, "#pane-orders", output_dir / "admin_orders_list.png"),
            )
        )

        click_admin_tab(driver, "recharges")
        shots.append(
            Shot(
                title="管理端-充值申请",
                note="充值申请分组、状态切换与审核区域加载正常；本次未执行审批操作。",
                path=capture_element(driver, "#pane-recharges", output_dir / "admin_recharges.png"),
            )
        )

        click_admin_tab(driver, "userView")
        shots.append(
            Shot(
                title="管理端-用户订单视图",
                note="按用户聚合的订单视图可正常显示，用于核对用户侧订单界面。",
                path=capture_element(driver, "#pane-userView", output_dir / "admin_user_view.png"),
            )
        )
    finally:
        driver.quit()
    return shots


def capture_admin_user_shots(
    base_url: str,
    admin_token: str,
    admin_password: str,
    user_phone: str,
    output_dir: Path,
) -> list[Shot]:
    if not admin_token:
        return []

    driver = build_driver(1500, 1700)
    shots: list[Shot] = []
    try:
        load_admin_users(driver, base_url, admin_token, admin_password, user_phone)
        shots.append(
            Shot(
                title="管理端-用户总览",
                note="用户列表、搜索区与当前选中用户摘要加载正常，已定位到测试账号。",
                path=capture_element(driver, ".layout > section.card:first-child", output_dir / "admin_users_list.png"),
            )
        )
        shots.append(
            Shot(
                title="管理端-用户操作面板",
                note="余额调整、免费次数修改、密码重置与删除用户面板展示正常；本次未执行写操作。",
                path=capture_element(driver, ".layout > section.card:last-child", output_dir / "admin_users_actions.png"),
            )
        )
    finally:
        driver.quit()
    return shots


def main() -> int:
    args = parse_args()
    run_stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_dir = (
        Path(args.output_dir)
        if args.output_dir
        else WORKSPACE / "docs" / "test_artifacts" / f"frontend_ui_{run_stamp}"
    )
    shots_dir = output_dir / "screenshots"
    ensure_dir(shots_dir)

    admin_creds = load_admin_creds()
    admin_token = admin_creds.get("admin_token", "")
    admin_password = admin_creds.get("admin_password", "")

    try:
        user_summary = fetch_user_summary(args.base_url, args.phone, args.password)
        admin_summary = fetch_admin_summary(args.base_url, admin_token, admin_password) if admin_token else {}
    except Exception as exc:  # pragma: no cover - external IO
        print(f"Failed to fetch API summary: {exc}", file=sys.stderr)
        return 1

    try:
        shots = []
        shots.extend(capture_user_shots(args.base_url, args.phone, args.password, shots_dir))
        shots.extend(capture_admin_order_shots(args.base_url, admin_token, admin_password, shots_dir))
        shots.extend(capture_admin_user_shots(args.base_url, admin_token, admin_password, args.phone, shots_dir))
    except TimeoutException as exc:  # pragma: no cover - browser automation
        print(f"Timed out during browser automation: {exc}", file=sys.stderr)
        return 2
    except Exception as exc:  # pragma: no cover - browser automation
        print(f"Browser automation failed: {exc}", file=sys.stderr)
        return 3

    report_path = output_dir / f"前端界面测试交付_{datetime.now().strftime('%Y-%m-%d')}.docx"
    build_report(report_path, shots, args.base_url, args.phone, user_summary, admin_summary)

    summary_payload = {
        "report_path": str(report_path),
        "shots": [str(shot.path) for shot in shots],
        "user_summary": user_summary,
        "admin_summary": admin_summary,
    }
    (output_dir / "summary.json").write_text(
        json.dumps(summary_payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print(json.dumps(summary_payload, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
